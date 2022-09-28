# Reading word documents

import docx

# calling docx.Document(), and passing the filename to open the .docx file.
# This will return a Document object, which has a paragraphs attribute that
# is a list of Paragraph objects.
my_doc = docx.Document("fileSamp.docx")

# check the number of paragraphs in the document
print(f"Number of paragraphs in the document: {len(my_doc.paragraphs)}.")


# the Paragraph object has a text attribute that contains a string of the text
# in that paragraph.
# print(my_doc.paragraphs[2].text)

# Each Paragraph object also has a runs attribute that is a list of Run objects.
# Run objects also have a text attribute, containing just the text in
# that particular run.
# print(f"Number of runs in paragraph 3: {len(my_doc.paragraphs[3].runs)}")
#
# print(my_doc.paragraphs[3].runs[0].text)
# print(my_doc.paragraphs[3].runs[1].text)
# print(my_doc.paragraphs[3].runs[3].text)

# getting full text from a .docx file
def get_text(file_name):
    """
    The get_text() function opens the Word document, loops over all the
    Paragraph objects in the paragraphs list, and then appends their text to the list
    in full_docx_txt. After the loop, the strings in full_docx_txt are joined
    together with newline characters.

    """

    # get the document obj
    doc = docx.Document(file_name)

    # list to hold the paragraphs
    full_docx_txt = []

    for paragraph in doc.paragraphs:
        full_docx_txt.append(paragraph.text)

    return "\n".join(full_docx_txt)


# print(get_text("fileSamp.docx"))

# writing word documents
# To create your own .docx file, call docx.Document() to return a new, blank
# Word Document object.
hello_docx = docx.Document()

# Calling add_heading() adds a paragraph with one of the heading styles.
# The arguments to add_heading() are a string of the heading text and
# an integer from 0 to 4. The integer 0 makes the heading the Title style,
# which is used for the top of the document.
hello_docx.add_heading("HELLO WORLD", 2)

# To add a line break (rather than starting a whole new paragraph), you can
# call the add_break() method on the Run object you want to have the break
# appear after.
hello_docx.paragraphs[0].runs[0].add_break()

# The add_paragraph() document method adds a new paragraph of text to the document
# and returns a reference to the Paragraph object that was added.
hello_docx.add_paragraph("Morbi velit neque, semper quis lorem quis, efficitur dignissim ipsum. Ut ac lorem sed "
                         "turpis imperdiet eleifend sit amet id sapien.")

# You can add paragraphs by calling the add_paragraph() method again
# with the new paragraph’s text. to add text to the end of an existing paragraph,
# you can call the paragraph’s add_run() method and pass it a string.
p_obj2 = hello_docx.add_paragraph("Vestibulum neque massa, scelerisque sit amet ligula eu, congue molestie mi. "
                                  "Praesent ut "
                                  "varius sem. Nullam at porttitor arcu, nec lacinia nisi. Ut ac dolor vitae odio "
                                  "interdum "
                                  "condimentum. Vivamus dapibus sodales ex, vitae malesuada ipsum cursus convallis. "
                                  "Maecenas "
                                  "sed egestas nulla, ac condimentum orci.")

p_obj2.add_run("•	Nulla facilisi.")

# Both add_paragraph() and add_run() accept an optional second argument
# that is a string of the Paragraph or Run object’s style.
p_obj3 = hello_docx.add_paragraph("This creates a new document from the built-in default template and saves it "
                                  "unchanged to a file named ‘test.docx’. The so-called “default template” is "
                                  "actually just a Word file having no content, stored with the installed python-docx "
                                  "package. ", "Title")


# passing a filename string to the save() document method to save the Document
# object to a file.
hello_docx.save("hello.docx")
